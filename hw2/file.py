from io import BytesIO
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
from PIL import Image

def base64_check(data):
    try:
        return base64.b64encode(base64.b64decode(data)) == data
    except Exception:
        return False

def convert_binary(data, into: str):
    if into == "PIL":
        if type(data) == str:
            return base64.b64decode(data.encode())
        elif not base64_check(data):  # bytes but not in base 64
            return data
        else:
            return base64.b64decode(data)
    elif into == "binary":
        if type(data) == str:
            return data.encode()
        elif not base64_check(data):  # bytes but not in base 64
            return base64.b64encode(data)
        else:
            return data
    elif into == "string":
        if type(data) == str:
            return data
        elif not base64_check(data):  # bytes but not in base 64
            data = base64.b64encode(data)
        return data.decode()
    else:
        return data

class File():
    file_wrapper: object
    file_destination: str
    file_name: str
    file_type: str

    def create_full_path(self):
        return self.file_destination + self.file_name + '.' + self.file_type

    def __init__(self, path: str, name: str, type: str):
        self.file_destination = path
        self.file_name = name
        self.file_type = type
        self.file_wrapper = self.create_instance()

    def create_instance(self):
        if self.file_type == "txt":
            return open(self.create_full_path(), 'w', encoding='utf-8')
        elif self.file_type == "html":
            file = open(self.create_full_path(), 'w', encoding='utf-8')
            file.write("<html>\n<head></head>\n<body>\n")
            return file
        elif self.file_type == "docx":
            return Document()

    def write_text(self, string: str):
        if self.file_type == "docx":
            self.file_wrapper.add_paragraph(string)
        elif self.file_type == "txt":
            self.file_wrapper.write(string + "\n")
        elif self.file_type == "html":
            self.file_wrapper.write("<p>" + string + "</p>" + "\n")

    def write_image(self, data):
        if self.file_type == "docx":
            try:
                new_file = BytesIO()
                Image.open(BytesIO(convert_binary(data, "PIL"))).convert('RGB').save(new_file, format="png")

                p = self.file_wrapper.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r = p.add_run()
                r.add_picture(new_file)  # , width=Cm(image_width))
            except:
                return